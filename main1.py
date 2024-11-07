import gradio as gr
from docx import Document
from deep_translator import GoogleTranslator
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_table_border(table):
    tbl = table._element
    tblPr = tbl.xpath("./w:tblPr")[0]
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')  # 1 px border size (size in half-points)
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')  # Black color
        tblBorders.append(border)
    tblPr.append(tblBorders)

def is_number(text):
    try:
        float(text)  # Check if the text can be converted to a float (handles decimals)
        return True
    except ValueError:
        return False

def translate_word_file(file_path):
    # Load the original Word document
    doc = Document(file_path)
    
    # Create a new document for the bilingual output
    translated_doc = Document()
    
    # Iterate over the paragraphs in the document
    for para in doc.paragraphs:
        if para.text.strip():  # Skip empty paragraphs
            # Translate the text from English to Hindi
            translated_text = GoogleTranslator(source='en', target='hi').translate(para.text)
            
            # Add the Hindi translation followed by the original English text
            translated_doc.add_paragraph(translated_text)
            translated_doc.add_paragraph(para.text)
    
    # Iterate over the tables in the document
    for table in doc.tables:
        translated_table = translated_doc.add_table(rows=0, cols=len(table.columns))
        set_table_border(translated_table)  # Set the table border to black

        for row in table.rows:
            translated_row = translated_table.add_row().cells
            for i, cell in enumerate(row.cells):
                if cell.text.strip():  # Skip empty cells
                    # Check if the cell text is numeric, including decimals
                    if is_number(cell.text):
                        # Add only the English number if it's numeric (integer or decimal)
                        translated_row[i].text = cell.text
                    else:
                        # Translate and add both Hindi and English content for non-numeric text
                        translated_text = GoogleTranslator(source='en', target='hi').translate(cell.text)
                        translated_row[i].text = f"{translated_text}\n{cell.text}"
                else:
                    translated_row[i].text = cell.text

    # Save the bilingual document
    output_path = "bilingual_output.docx"
    translated_doc.save(output_path)

    return output_path

# Create the Gradio interface
interface = gr.Interface(
    fn=translate_word_file,
    inputs=gr.File(type="filepath", label="Upload English Word File"),
    outputs=gr.File(label="Download Bilingual Word File"),
    title="English to Hindi Word Translator",
    description="Upload an English Word file to translate its contents into Hindi. The output will be a bilingual document with Hindi followed by English."
)

# Launch the Gradio app
interface.launch()
