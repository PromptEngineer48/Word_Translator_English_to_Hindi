import gradio as gr
from docx import Document
from deep_translator import GoogleTranslator

def translate_word_file(file_path):
    # Load the English Word file
    doc = Document(file_path)
    
    # Create a new document for the bilingual translation
    translated_doc = Document()

    # Iterate over each paragraph and translate the text
    for para in doc.paragraphs:
        if para.text.strip():  # Skip empty paragraphs
            # Translate the text from English to Hindi
            translated_text = GoogleTranslator(source='en', target='hi').translate(para.text)
            
            # Add the Hindi translation and the original English text to the new document
            translated_doc.add_paragraph(translated_text)
            translated_doc.add_paragraph(para.text)

    # Save the bilingual document to a new Word file
    output_path = "bilingual_output.docx"
    translated_doc.save(output_path)

    return output_path

# Create the Gradio interface
interface = gr.Interface(
    fn=translate_word_file,
    inputs=gr.File(type="filepath", label="Upload English Word File"),
    outputs=gr.File(label="Download Bilingual Word File"),
    title="English to Hindi Word Translator",
    description="Upload an English Word file to translate its contents into Hindi. The output will be a bilingual document with Hindi followed by English."
)

# Launch the Gradio app
interface.launch()
